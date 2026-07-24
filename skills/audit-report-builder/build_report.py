#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
調查報告自動彙整程式 — skill 版本
可從命令列指定：
  --plan       調查計畫檔案路徑
  --workpapers 工作底稿檔案路徑（逗號分隔或多個 --workpapers）
  --template   樣板檔案路徑（預設 skill 目錄下 template/調查報告.docx）
  --output     輸出檔案路徑（預設 ./調查報告_彙整完成.docx）
"""

import os, re, sys, json
from collections import OrderedDict

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

try:
    import win32com.client
    HAS_WIN32COM = True
except ImportError:
    HAS_WIN32COM = False
    print("警告：未安裝 pywin32，無法讀取 .doc 格式")
    print("請執行：pip install pywin32")

ZH_NUMS = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]

def zh_num(n):
    if 1 <= n <= 10:
        return ZH_NUMS[n]
    if 11 <= n <= 19:
        return "十" + (ZH_NUMS[n - 10] if n > 10 else "")
    if 20 <= n <= 29:
        return "二十" + (ZH_NUMS[n - 20] if n > 20 else "")
    return str(n)

def xml_escape(text):
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def _set_indent_xml(p, left_twips, first_line_twips):
    pPr = p._element.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    hanging_attr = qn('w:hanging')
    if hanging_attr in ind.attrib:
        del ind.attrib[hanging_attr]
    ind.set(qn('w:left'), str(int(round(left_twips))))
    ind.set(qn('w:firstLine'), str(int(round(first_line_twips))))

# ================================================================
#  1. 讀取 .doc / .docx 檔案
# ================================================================

def read_doc(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.doc' and HAS_WIN32COM:
        return _read_doc_win32(filepath)
    elif ext == '.docx':
        return _read_docx(filepath)
    else:
        raise ValueError(f"不支援的檔案格式：{ext}")

def _read_doc_win32(filepath):
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(filepath))
        paragraphs = []
        for p in doc.Paragraphs:
            txt = p.Range.Text.strip()
            style = p.Style.NameLocal if p.Style else ""
            paragraphs.append({"text": txt, "style": style})
        tables = []
        for ti, tbl in enumerate(doc.Tables):
            rows = []
            for ri in range(tbl.Rows.Count):
                row = []
                for ci in range(tbl.Columns.Count):
                    try:
                        cell = tbl.Cell(ri + 1, ci + 1)
                        txt = cell.Range.Text.strip().replace('\r\x07', '').replace('\x07', '')
                        row.append(txt)
                    except Exception:
                        row.append("")
                rows.append(row)
            caption_before = ""
            caption_after = ""
            # try to find caption before/after
            try:
                rng = tbl.Range
                rng.Start = rng.Start - 200
                rng.End = tbl.Range.Start
                prefix = rng.Text.strip()
                lines = prefix.split('\r')
                for line in lines:
                    ls = line.strip()
                    if re.match(r'^[表圖]\s*\d', ls):
                        caption_before = ls
            except Exception:
                pass
            try:
                rng = tbl.Range
                rng.Start = tbl.Range.End
                rng.End = rng.End + 200
                suffix = rng.Text.strip()
                lines = suffix.split('\r')
                for line in lines:
                    ls = line.strip()
                    if re.search(r'資料來源', ls):
                        caption_after = ls
            except Exception:
                pass
            tables.append({
                "rows": rows,
                "caption_before": caption_before,
                "caption_after": caption_after,
                "has_header_shading": True,
            })
        doc.Close()
        return {"paragraphs": paragraphs, "tables": tables}
    finally:
        word.Quit()

def _read_docx(filepath):
    d = Document(filepath)
    paragraphs = []
    for p in d.paragraphs:
        txt = p.text.strip()
        style = p.style.name if p.style else ""
        paragraphs.append({"text": txt, "style": style})
    tables = []
    for tbl in d.tables:
        rows = []
        for ri, row in enumerate(tbl.rows):
            r = [cell.text.strip().replace('\x07', '') for cell in row.cells]
            rows.append(r)
        tables.append({
            "rows": rows,
            "caption_before": "",
            "caption_after": "",
            "has_header_shading": True,
        })
    return {"paragraphs": paragraphs, "tables": tables}

# ================================================================
#  2. 分析調查計畫
# ================================================================

PLAN_KEYWORDS = [
    ("查核緣起", ["壹", "一", "調查緣起", "查核緣起"]),
    ("查核依據", ["貳", "二", "調查依據", "查核依據"]),
    ("查核範圍", ["參", "三", "調查範圍", "查核範圍"]),
    ("查核過程", ["肆", "四", "調查過程", "查核過程"]),
    ("查核重點", ["伍", "五", "調查重點", "查核重點"]),
]

# 欲拋棄的法規附件關鍵字
LAW_KEYWORDS = ["政府採購法", "行政程序法", "會計法", "審計法", "機關主會計",
                "政府採購", "預算法", "各機關", "政府支出", "內部審核",
                "機關辦理", "中央政府", "審計部", "審計機關", "各機關單位",
                "行政院", "政府會計", "審計準則"]

WP_HEADINGS = [
    "採購規範與預算編列",
    "採購決標、履約執行情形",
    "採購決標、履約執行",
    "工程使用情形",
    "工程驗收與財產登錄",
    "其他調查事項",
]

def clean_text(text):
    if not text:
        return ""
    text = text.replace('\r', '').replace('\x07', '').replace('\u3000', '')
    chars = []
    for ch in text:
        if ord(ch) < 32 and ch not in ('\t', '\n'):
            continue
        chars.append(ch)
    result = ''.join(chars).strip()
    result = re.sub(r' {2,}', ' ', result)
    return result

def is_law_text(text):
    return any(kw in text for kw in LAW_KEYWORDS)

def analyze_plan(data):
    """分析調查計畫：將段落分配到各章節"""
    sections = OrderedDict()
    for name, _ in PLAN_KEYWORDS:
        sections[name] = []

    current = None
    plan_paras = data["paragraphs"]
    report_title = plan_paras[0]["text"].strip() if plan_paras else ""
    if len(report_title) > 20:
        pass
    elif len(plan_paras) > 1:
        report_title = plan_paras[1]["text"].strip()
    plan_title = True

    for p in data["paragraphs"]:
        txt = p["text"].strip()
        if plan_title:
            plan_title = False
            continue

        matched = None
        for sec_name, keywords in PLAN_KEYWORDS:
            for kw in keywords:
                num_kw = r'^[壹貳參肆伍]+\s*[、．]\s*' + re.escape(kw) + r'$'
                num_kw2 = r'^[一二三四五]+\s*[、．]\s*' + re.escape(kw) + r'$'
                if re.match(num_kw, txt) or re.match(num_kw2, txt):
                    matched = sec_name
                    break
                if txt == kw:
                    matched = sec_name
                    break
            if matched:
                break

        if matched:
            current = matched
            continue

        skip_patterns = [
            r'^第[一二三四五六七八九十]+節',
            r'^[一二三四五六七八九十]+[、．]',
            r'^附註',
            r'^資料來源',
            r'^表\d',
            r'^(圖|表)\s+\d',
        ]
        is_skip = any(re.match(sp, txt) for sp in skip_patterns)
        if is_skip:
            current = None
            continue

        law_header_keywords = ["參考法令規章", "參考法規", "相關法規"]
        if any(txt == kw for kw in law_header_keywords):
            current = None
            continue

        if current and current in sections:
            sections[current].append({"text": txt, "style": p.get("style", "")})

    # 將「調查對象」「調查範圍」抽離自成一組
    relocate_keys = ["調查對象", "調查範圍"]
    scope_items = []
    for sec_name in ["查核依據", "查核緣起"]:
        kept = []
        for p in sections.get(sec_name, []):
            if any(p["text"].startswith(k + "：") for k in relocate_keys):
                scope_items.append(p)
            else:
                kept.append(p)
        sections[sec_name] = kept

    plan = {
        "title": report_title,
        "sections": sections,
        "tables": data.get("tables", []),
        "scope_items": scope_items,
        "normal_paras": sections.get("查核緣起", []),
    }
    return plan

# ================================================================
#  3. 分析工作底稿
# ================================================================

def detect_wp_heading(text):
    for h in WP_HEADINGS:
        if text.strip() == h:
            return h
    m = re.match(r'^[（(][一二三四五六七八九十]+[）)]\s*', text)
    if m:
        return None  # level 2 numbering: continue
    return None

def is_skip_para(text):
    """過濾結構性文字，不納入查核事實"""
    t = text.strip()
    if not t:
        return True
    skips = [
        r'^聯絡人', r'^附件資料', r'^資料來源', r'^本案由', r'^本項由',
        r'^單位[：:]', r'^擬議', r'^案件基本資料', r'^主辦', r'^輔助',
        r'^電話', r'^傳真', r'^表\d+', r'^其他相關調查事項',
        r'宜蘭縣[\S]*部分', r'宜蘭縣[\S]*公所$',
    ]
    for pat in skips:
        if re.match(pat, t):
            return True
    return False

def is_table_caption(text):
    return bool(re.match(r'^(表|圖)\s*\d', text.strip()))

def analyze_workpaper(data, idx):
    paragraphs = data["paragraphs"]
    tables = data.get("tables", [])

    # 提取鄉鎮名稱
    township = ""
    for p in paragraphs:
        m = re.match(r'^(宜蘭縣\S{2,6}部分)', p["text"].strip())
        if m:
            township = m.group(1)
            break

    # 偵測大綱 heading 並分類段落
    sections = OrderedDict()
    current_section = None
    subsection_texts = []
    section_heading = None

    for p in paragraphs:
        txt = clean_text(p["text"])
        if not txt:
            continue

        # 跳過結構性文字
        if is_skip_para(txt):
            continue
        if is_table_caption(txt):
            continue

        # 檢查是否為 WP_HEADING
        h = detect_wp_heading(txt)
        if h:
            current_section = h
            sections.setdefault(current_section, [])
            section_heading = h
            continue

        # 檢查是否為子段落 heading（如注意事項、建議處理意見等）
        sub_h = txt.strip()
        if sub_h in ("注意事項", "建議處理意見", "查核事實概述",
                      "待查證事項", "待查核事項", "查證事項",
                      "工程使用情形", "其他調查事項"):
            current_section = sub_h
            sections.setdefault(current_section, [])
            section_heading = sub_h
            continue
        # 也檢查原文可能會有的小標題
        m = re.match(r'^[（(][一二三四五六七八九十]+[）)]\s*\S', txt)
        if m and current_section:
            sections.setdefault(current_section, [])
            sections[current_section].append({"text": txt, "format": "numbered"})
            continue

        if current_section:
            sections.setdefault(current_section, [])
            sections[current_section].append({"text": txt, "format": "text"})
        else:
            # pre-section content
            sections.setdefault("前言", [])
            sections["前言"].append({"text": txt, "format": "text"})

    # 為表格分類 section
    for tbl in tables:
        cap = (tbl.get("caption_before", "") or "")
        if "預算" in cap or "規範" in cap or "編列" in cap:
            tbl["section"] = "採購規範與預算編列"
        elif "決標" in cap or "履約" in cap or "執行" in cap:
            tbl["section"] = "採購決標、履約執行情形"
        elif "使用" in cap or "減損" in cap:
            tbl["section"] = "工程使用情形"
        elif "驗收" in cap or "財產" in cap or "登錄" in cap:
            tbl["section"] = "工程驗收與財產登錄"
        elif "待查證" in cap or "詢證" in cap:
            tbl["section"] = "待查證事項"
        else:
            tbl["section"] = "其他調查事項"

        # 填補 caption_before/after
        if not tbl.get("caption_before"):
            tbl["caption_before"] = ""
        if not tbl.get("caption_after"):
            tbl["caption_after"] = ""

    result = {
        "township": township,
        "sections": sections,
        "tables": tables,
        "index": idx,
    }
    return result

# ================================================================
#  4. 表格處理
# ================================================================

def make_table_xml(rows_data, has_header_shading):
    if not rows_data:
        return None
    n_rows = len(rows_data)
    n_cols = max(len(r) for r in rows_data)
    default_w = max(800, 5000 // n_cols)
    grid = '<w:tblGrid>'
    for _ in range(n_cols):
        grid += f'<w:gridCol w:w="{default_w}"/>'
    grid += '</w:tblGrid>'
    border = (
        '<w:tblBorders>'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    rows_xml = ''
    for ri, row in enumerate(rows_data):
        cells_xml = ''
        for ci in range(n_cols):
            txt = xml_escape(row[ci] if ci < len(row) else "")
            shading = ''
            if ri == 0 and has_header_shading:
                shading = '<w:shd w:val="clear" w:color="auto" w:fill="E7E6E6"/>'
            is_bold = '1' if (ri == 0) else '0'
            lines = txt.split('\n')
            paras = ''
            for line in lines:
                b = '<w:b/>' if is_bold == '1' else ''
                paras += (
                    f'<w:p><w:pPr><w:jc w:val="center"/></w:pPr>'
                    f'<w:r><w:rPr>'
                    f'<w:rFonts w:eastAsia="\u6a19\u6977\u9ad4"/>'
                    f'<w:sz w:val="24"/><w:szCs w:val="24"/>'
                    f'{b}'
                    f'</w:rPr><w:t xml:space="preserve">{line}</w:t></w:r></w:p>'
                )
            cells_xml += (
                f'<w:tc>'
                f'<w:tcPr><w:vAlign w:val="center"/>{shading}</w:tcPr>'
                f'{paras}'
                f'</w:tc>'
            )
        rows_xml += f'<w:tr>{cells_xml}</w:tr>'
    xml = (
        f'<w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        f'<w:tblPr>'
        f'<w:tblW w:w="0" w:type="auto"/>'
        f'<w:jc w:val="center"/>'
        f'{border}'
        f'</w:tblPr>'
        f'{grid}'
        f'{rows_xml}'
        f'</w:tbl>'
    )
    return parse_xml(xml)

def table_fingerprint(tbl):
    rows = tbl.get("rows", [])
    fp = tuple(tuple(c[:20] for c in row) for row in rows[:5])
    return fp

# ================================================================
#  5. 建立調查報告
# ================================================================

GLOBAL_TABLE_NUM = [0]

CHINESE_NUM = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
               "十一", "十二", "十三", "十四", "十五"]

def _to_fullwidth(text):
    return text.replace('(', '（').replace(')', '）')

# 字元縮排對照（1字元 = 16pt = 320twips）
def _char_to_twips(n):
    return int(round(n * 16 * 20))

def _detect_level(text):
    m = re.match(r'^[一二三四五六七八九十]+[.．、]', text)
    if m:
        return ("L2", 1, -2)
    m = re.match(r'^[（(][一二三四五六七八九十]+[）)]', text)
    if m:
        return ("L3", 3, -2)
    m = re.match(r'^\d+\s*[.．、]', text)
    if m:
        return ("L4", 6, -2)
    return None

def add_para(doc, text, style_name='Body Text Indent 2', number_prefix=None):
    if not text:
        return
    p = doc.add_paragraph()
    p.style = doc.styles[style_name]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    final_text = _to_fullwidth((number_prefix + text) if number_prefix else text)
    level = _detect_level(final_text)
    if level:
        lv, lc, fc = level
        _set_indent_xml(p, _char_to_twips(lc), _char_to_twips(fc))
    else:
        _set_indent_xml(p, _char_to_twips(4), _char_to_twips(2))
    run = p.add_run(final_text)
    run.bold = level is not None
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(16)

def _numbered_para(doc, text, left_twips, first_line_twips):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_indent_xml(p, left_twips, first_line_twips)
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    run = p.add_run(_to_fullwidth(text))
    run.bold = True
    run.font.name = '標楷體'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    run.font.size = Pt(16)

def add_numbered_paras(doc, paras):
    if not paras:
        return
    texts = [p["text"] for p in paras]
    if len(texts) == 1:
        add_para(doc, texts[0])
        return
    level1_idx = 0
    for t in texts:
        m2 = re.match(r'^[（(][一二三四五六七八九十]+[）)]', t)
        m3 = re.match(r'^\d+\s*[.．、]', t)
        if m2 or m3:
            add_para(doc, t)
        else:
            level1_idx += 1
            prefix = f"{CHINESE_NUM[level1_idx-1]}、" if level1_idx <= len(CHINESE_NUM) else f"{level1_idx}、"
            add_para(doc, prefix + t)

def add_caption(doc, text, num):
    if not text:
        return
    cleaned = re.sub(r'^(表|第)\s*\d+\s*', '', text).strip()
    if not cleaned:
        return
    is_note = re.match(r'^(單位[：:]|註[：:]|註記[：:]|資料來源[：:])', cleaned)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption = _to_fullwidth(f'表{num}  {cleaned}')
    run = p.add_run(caption)
    run.bold = False
    run.font.name = '標楷體'
    run.font.size = Pt(10) if is_note else Pt(14)

def add_source(doc, text):
    if not text:
        return
    cleaned = re.sub(r'^資料來源[：:]?\s*', '', text).strip()
    if not cleaned:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.left_indent = Cm(0)
    run = p.add_run(_to_fullwidth(f'資料來源：{cleaned}'))
    run.font.name = '標楷體'
    run.font.size = Pt(10)

def insert_table(doc, tbl_data):
    global GLOBAL_TABLE_NUM
    GLOBAL_TABLE_NUM[0] += 1
    num = GLOBAL_TABLE_NUM[0]
    rows = tbl_data.get("rows", [])
    cap_before = tbl_data.get("caption_before", "")
    cap_after = tbl_data.get("caption_after", "")
    has_shading = tbl_data.get("has_header_shading", True)

    if cap_before and not re.match(r'^表\d+$', cap_before.strip()):
        add_caption(doc, cap_before, num)
    tbl_elem = make_table_xml(rows, has_shading)
    if tbl_elem is not None:
        doc.element.body.append(tbl_elem)
    if cap_after:
        add_source(doc, cap_after)

def _extract_scope_key(text):
    for key in ("調查對象", "調查範圍"):
        if text.startswith(key + "："):
            return key
    return None

def _add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(_to_fullwidth(text))
    r.bold = True
    r.font.name = '標楷體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    r.font.size = Pt(16)

def _add_level1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_indent_xml(p, _char_to_twips(1), _char_to_twips(-2))
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    r = p.add_run(_to_fullwidth(text))
    r.bold = True
    r.font.name = '標楷體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    r.font.size = Pt(16)

def build_report(template_path, output_path, plan, wp_list):
    doc = Document(template_path)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    if sectPr is not None:
        body.remove(sectPr)
    children = list(body)
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            body.remove(child)

    REPORT_TITLE = _to_fullwidth(plan.get("title", "") or "調查報告")
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.line_spacing = Pt(28)
    title_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    title_p.paragraph_format.space_after = Pt(18)
    r = title_p.add_run(REPORT_TITLE)
    r.bold = True
    r.font.name = '標楷體'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    r.font.size = Pt(18)

    GLOBAL_TABLE_NUM[0] = 0
    inserted_tables = set()

    # ======== 壹、查核緣起 ========
    _add_heading(doc, "壹、查核緣起")
    scope_items = plan.get("scope_items", [])
    scope_texts = {_extract_scope_key(p["text"]): p["text"] for p in scope_items}
    normal_paras = [p for p in plan["sections"].get("查核緣起", [])
                    if _extract_scope_key(p["text"]) not in scope_texts]
    add_numbered_paras(doc, normal_paras)
    if scope_items:
        _add_level1(doc, "一、調查對象及範圍")
        for i, p in enumerate(scope_items):
            prefix = f"({CHINESE_NUM[i]})"
            _numbered_para(doc, prefix + p["text"], _char_to_twips(3), _char_to_twips(-2))

    # ======== 貳、查核依據 ========
    _add_heading(doc, "貳、查核依據")
    basis_paras = [p for p in plan["sections"].get("查核依據", [])
                   if "號函辦理" in p["text"]]
    add_numbered_paras(doc, basis_paras)

    # ======== 參、查核範圍 ========
    _add_heading(doc, "參、查核範圍")
    add_numbered_paras(doc, plan["sections"].get("查核範圍", []))
    for tbl in plan.get("tables", []):
        insert_table(doc, tbl)

    # ======== 肆、查核限制 ========
    _add_heading(doc, "肆、查核限制")
    add_para(doc, f"本次查核囿於審計人力及查核時間，係依調查計畫所列各項查核重點，抽查機關辦理「{REPORT_TITLE}」所提供之書面資料，恐無法發現所有缺失，亦無法確保無錯誤或舞弊之情事發生。")

    # ======== 伍、查核重點 ========
    _add_heading(doc, "伍、查核重點")
    focus_paras = []
    for p in plan["sections"].get("查核重點", []):
        focus_paras.append(p)
        if "其他相關調查事項" in p["text"]:
            break
    add_numbered_paras(doc, focus_paras)

    # ======== 陸、查核事實 ========
    _add_heading(doc, "陸、查核事實")
    seen_text = set()
    all_facts = []
    for wp in wp_list:
        for sname in ["查核事實概述"] + WP_HEADINGS:
            for pdata in wp["sections"].get(sname, []):
                t = pdata["text"].strip()
                if not t:
                    continue
                if is_skip_para(t):
                    continue
                if is_table_caption(t):
                    continue
                fp = t[:40]
                if fp in seen_text:
                    continue
                seen_text.add(fp)
                all_facts.append(t)
    for t in all_facts:
        add_para(doc, t)

    # 陸表格：查核事實相關表格（排除待查證及意見類）
    for wp in wp_list:
        for tbl in wp["tables"]:
            sec = tbl.get("section", "")
            if sec in ("待查證事項", "其他調查事項"):
                continue
            fp = table_fingerprint(tbl)
            if fp not in inserted_tables:
                insert_table(doc, tbl)
                inserted_tables.add(fp)

    # ======== 柒、查核意見 ========
    _add_heading(doc, "柒、查核意見")
    for wp in wp_list:
        township = wp.get("township", "")
        if not township:
            continue
        wlabel = f"〔工作底稿（{wp['index']}）〕"
        _numbered_para(doc, f"{township}{wlabel}", _char_to_twips(1), _char_to_twips(-2))

        # 注意事項子標題
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(1.1)
        p.paragraph_format.first_line_indent = Cm(0.6)
        p.paragraph_format.line_spacing = Pt(28)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        r = p.add_run("注意事項")
        r.bold = True
        r.font.name = '標楷體'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
        r.font.size = Pt(16)

        content_paras = []
        for sec in ("注意事項", "建議處理意見"):
            for pdata in wp["sections"].get(sec, []):
                t = pdata["text"].strip()
                if not t or t in ("注意事項", "建議處理意見", "查核事實概述"):
                    continue
                if re.match(r'^(表|圖)\s*\d', t):
                    continue
                if re.search(r'資料來源|自行整理|單位[：:]|新臺幣', t):
                    continue
                if "：" not in t and len(t) < 8:
                    continue
                content_paras.append((sec, t))

        is_first_item = True
        for sec, t in content_paras:
            if re.search(r'函請.*部分', t):
                _numbered_para(doc, t, _char_to_twips(1), _char_to_twips(-2))
                is_first_item = True
                continue
            if re.match(r'^[（(]?[一二三四五六七八九十]+[）)]?\s*[建查]', t) or (sec in ("建議處理意見",) and not re.match(r'^[（(]', t)):
                _numbered_para(doc, "（一）" + t if not re.match(r'^[（(]', t) else t, _char_to_twips(3), _char_to_twips(-2))
                is_first_item = True
                continue
            pp = doc.add_paragraph()
            pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pp.paragraph_format.left_indent = Cm(1.7)
            if is_first_item:
                pp.paragraph_format.first_line_indent = Cm(0)
            else:
                pp.paragraph_format.first_line_indent = Pt(32)
            pp.paragraph_format.line_spacing = Pt(28)
            pp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            rr = pp.add_run(_to_fullwidth(t))
            rr.bold = is_first_item
            rr.font.name = '標楷體'
            rr._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
            rr.font.size = Pt(16)
            is_first_item = False

        for tbl in wp["tables"]:
            if tbl.get("section") in ("待查證事項", "其他調查事項"):
                fp = table_fingerprint(tbl)
                if fp not in inserted_tables:
                    insert_table(doc, tbl)
                    inserted_tables.add(fp)

    # ======== 捌、擬議處理意見 ========
    _add_heading(doc, "捌、擬議處理意見")
    add_para(doc, "上開查核意見，擬分別通知南澳鄉及大同鄉公所注意檢討改善，並將調查結果（調查報告及審核通知）函送第五廳彙辦，電子檔案另送至第五廳承辦人員電子郵件信箱（phuang@mail.audit.gov.tw）。")

    # ======== 玖、附錄 ========
    _add_heading(doc, "玖、附錄")
    add_para(doc, "一、抽查紀錄（工作底稿）")
    for i in range(1, len(wp_list) + 1):
        add_para(doc, f"（{zh_num(i)}）工作底稿（{i}）")
    add_para(doc, "二、調查計畫")

    if sectPr is not None:
        body.append(sectPr)
    doc.save(output_path)
    print(f"報告已成功產出：{output_path}")

# ================================================================
#  6. skill 入口：處理使用者提供的檔案路徑或上傳檔案
# ================================================================

def main():
    import argparse
    parser = argparse.ArgumentParser(description="調查報告彙整工具")
    parser.add_argument("--plan", required=True, help="調查計畫檔案路徑（.doc/.docx）")
    parser.add_argument("--workpapers", required=True, action="append",
                        help="工作底稿檔案路徑（可重複使用）")
    parser.add_argument("--template", default=None,
                        help="樣板檔路徑（預設使用內建樣板）")
    parser.add_argument("--output", default=None, help="輸出檔路徑")
    args = parser.parse_args()

    # 樣板與輸出路徑
    skill_dir = os.path.dirname(os.path.abspath(__file__))
    template = args.template or os.path.join(skill_dir, "template", "調查報告.docx")
    if not os.path.exists(template):
        print(f"ERR: 找不到樣板 {template}")
        sys.exit(1)

    output = args.output or os.path.join(os.getcwd(), "調查報告_彙整完成.docx")

    print("=" * 60)
    print("調查報告彙整工具")
    print("=" * 60)

    # 讀取調查計畫
    print("\n[1/4] 讀取調查計畫...")
    if not os.path.exists(args.plan):
        print(f"  ERR: 找不到 {args.plan}")
        sys.exit(1)
    try:
        raw = read_doc(args.plan)
        plan = analyze_plan(raw)
        print(f"  OK: {len(raw['paragraphs'])} 段落, {len(raw['tables'])} 表格")
    except Exception as e:
        print(f"  ERR: {e}")
        import traceback; traceback.print_exc()
        sys.exit(1)

    # 讀取工作底稿
    print("\n[2/4] 讀取工作底稿...")
    wp_list = []
    for i, wp_path in enumerate(args.workpapers, 1):
        if not os.path.exists(wp_path):
            print(f"  ERR: 找不到 {wp_path}")
            continue
        try:
            raw = read_doc(wp_path)
            wp = analyze_workpaper(raw, i)
            wp_list.append(wp)
            print(f"  OK: {os.path.basename(wp_path)} - {len(raw['paragraphs'])} 段落, {len(raw['tables'])} 表格")
        except Exception as e:
            print(f"  ERR: {os.path.basename(wp_path)} - {e}")
            import traceback; traceback.print_exc()

    if not wp_list:
        print("  [!] 未找到任何工作底稿，仍可產出報告（僅含調查計畫內容）")

    # 建立報告
    print("\n[3/4] 建立調查報告...")
    try:
        build_report(template, output, plan, wp_list)
        print("[4/4] 完成！")
    except Exception as e:
        print(f"  ERR: {e}")
        import traceback; traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()
